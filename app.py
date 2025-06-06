import os
from flask import Flask, request, redirect, url_for, render_template, session, send_from_directory
from werkzeug.utils import secure_filename
import fitz # PyMuPDF
import anthropic
from docx import Document
from datetime import datetime
import shutil # For copying placeholder image
import google.generativeai as genai # For Gemini API
import openai # For OpenAI API
import requests # For fetching image from URL if Gemini returns that
from PIL import Image # For image processing if Gemini returns image bytes
import io # For image processing if Gemini returns image bytes


UPLOAD_FOLDER = 'uploads'
DOCS_FOLDER = 'docs' # For saving summaries
# Define DOCS_DIR for send_from_directory
DOCS_DIR = os.path.abspath(DOCS_FOLDER)

STATIC_IMAGES_FOLDER = 'static/images' # For saving visualizations

ALLOWED_EXTENSIONS = {'pdf'}

app = Flask(__name__)

# --- API Key Configuration ---
# User must set these as environment variables for the application to function with live AI services.
# For example, in your terminal before running the app:
# export OPENAI_API_KEY='your_openai_key_here'
# export ANTHROPIC_API_KEY='your_anthropic_key_here'
# export GEMINI_API_KEY='your_gemini_key_here'
app.config['OPENAI_API_KEY'] = os.environ.get('OPENAI_API_KEY')
app.config['ANTHROPIC_API_KEY'] = os.environ.get('ANTHROPIC_API_KEY')
app.config['GEMINI_API_KEY'] = os.environ.get('GEMINI_API_KEY')
# --- End API Key Configuration ---

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'super secret key'  # Needed for session management

# Create upload, docs, and static/images folders if they don't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
os.makedirs(DOCS_FOLDER, exist_ok=True) # DOCS_DIR is derived from this
# app.static_folder is 'static' by default. We want 'static/images'
os.makedirs(os.path.join(app.static_folder, 'images'), exist_ok=True)


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def parse_structured_summary(summary_text):
    """Parses a string summary into a dictionary by section headers."""
    parsed = {}
    current_section = None
    current_content = []

    if not isinstance(summary_text, str): # Handle cases where summary might not be a string
        return {"Error": "Summary was not in expected string format."}

    for line in summary_text.strip().split('\n'):
        line_stripped = line.strip()
        if line_stripped.startswith("Abstract:") or \
           line_stripped.startswith("Introduction:") or \
           line_stripped.startswith("Results:") or \
           line_stripped.startswith("Discussion:"):
            if current_section and current_content:
                parsed[current_section] = " ".join(current_content).strip()

            parts = line_stripped.split(":", 1)
            current_section = parts[0].strip()
            current_content = [parts[1].strip()] if len(parts) > 1 else []
        elif current_section:
            current_content.append(line_stripped)

    if current_section and current_content: # Add the last section
        parsed[current_section] = " ".join(current_content).strip()

    if not parsed: # If no sections were found, return the original text under a generic key
        return {"Full Summary": summary_text}

    return parsed

def generate_visualization_prompt_with_anthropic(summary_text_or_dict):
    """
    Generates a visualization prompt using Anthropic API based on the summary.
    Accepts either a string summary or a dictionary of parsed summary sections.
    """
    anthropic_api_key = app.config.get('ANTHROPIC_API_KEY')
    if not anthropic_api_key:
        print("ANTHROPIC_API_KEY not found. Returning default visualization prompt.")
        return "Anthropic API key not set. Using default prompt: A generic scientific concept representing research findings."

    summary_input_string = ""
    if isinstance(summary_text_or_dict, dict):
        # Convert dict to a readable string format for the prompt
        for key, value in summary_text_or_dict.items():
            summary_input_string += f"{key}: {value}\n"
        summary_input_string = summary_input_string.strip()
    elif isinstance(summary_text_or_dict, str):
        summary_input_string = summary_text_or_dict
    else:
        print("Invalid summary format for Anthropic prompt generation. Using generic prompt.")
        return "Invalid summary input. Default prompt: Key scientific breakthroughs illustrated."

    if not summary_input_string.strip(): # Handle empty summary
        print("Empty summary provided for Anthropic prompt generation. Using generic prompt.")
        return "Empty summary. Default prompt: Abstract visualization of data."

    try:
        client = anthropic.Anthropic(api_key=anthropic_api_key)

        claude_prompt_text = f"""
Based on the following research paper summary, generate a concise and visually descriptive prompt (around 20-30 words) that can be used by an image generation AI to create a compelling visualization representing the core findings or essence of the paper.

Summary:
{summary_input_string}

Generated Image Prompt:
"""
        # The "Generated Image Prompt:" line helps instruct Claude to only return the prompt.

        response = client.messages.create(
            model="claude-3-sonnet-20240229", # Using Sonnet as a balance
            max_tokens=100,
            messages=[
                {"role": "user", "content": claude_prompt_text}
            ]
        )

        if response.content and response.content[0].text:
            visualization_prompt = response.content[0].text.strip()
            # Sometimes Claude might still add a preamble like "Here's a prompt:"
            # A simple way to try and remove it:
            if "Generated Image Prompt:" in visualization_prompt:
                 visualization_prompt = visualization_prompt.split("Generated Image Prompt:")[-1].strip()
            if visualization_prompt.lower().startswith("here's a prompt:") or visualization_prompt.lower().startswith("here is a prompt:"):
                visualization_prompt = visualization_prompt.split(":",1)[-1].strip()
            return visualization_prompt
        else:
            print("Anthropic API returned an empty or invalid response for visualization prompt.")
            return "Error: Anthropic API (Claude) returned an empty response for visualization prompt."

    except anthropic.APIConnectionError as e:
        print(f"Anthropic API Connection Error: {e}")
        return f"Error: Anthropic API Connection Error - {e}"
    except anthropic.RateLimitError as e:
        print(f"Anthropic API Rate Limit Exceeded: {e}")
        return f"Error: Anthropic API Rate Limit Exceeded - {e}"
    except anthropic.APIStatusError as e:
        print(f"Anthropic API Status Error: {e}")
        return f"Error: Anthropic API Status Error - {e}"
    except Exception as e:
        print(f"An unexpected error occurred with Anthropic: {e}")
        return f"Error: An unexpected error occurred with the Anthropic API - {e}"

def generate_image_with_ai(prompt, base_filename_prefix="visualization"):
    """
    Generates an image using an AI model (Gemini placeholder).
    Attempts to use Gemini API if key is present, otherwise copies a placeholder image.
    IMPORTANT NOTE: Direct text-to-image generation with the `google-generativeai`
    library is complex and model-dependent. As of early 2024, this library is more
    focused on text, chat, and multimodal understanding. True text-to-image generation
    (like with Google's Imagen models) is often accessed via the Vertex AI SDK or
    more specific APIs. The following implementation includes a conceptual structure
    for a Gemini API call but will likely default to placeholder logic due to these
    API specificities. A robust solution would require using the appropriate SDK and model name.
    """
    gemini_api_key = app.config.get('GEMINI_API_KEY')

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    original_filename_no_ext = os.path.splitext(base_filename_prefix)[0]
    image_filename = f"{original_filename_no_ext}_visualization_{timestamp}.png"
    image_save_subpath = os.path.join('images', image_filename) # Relative path for web
    image_abs_save_path = os.path.join(app.static_folder, image_save_subpath) # Absolute path for saving

    def _copy_placeholder_image():
        placeholder_src = os.path.join(app.static_folder, 'images', 'placeholder.png')
        if not os.path.exists(placeholder_src):
            print(f"Error: Placeholder image not found at {placeholder_src}")
            return None
        try:
            shutil.copy(placeholder_src, image_abs_save_path)
            print(f"Used placeholder image: {image_filename}")
            return image_save_subpath # Return web accessible path
        except Exception as e_copy:
            print(f"Error copying placeholder image: {e_copy}")
            return None

    if not gemini_api_key:
        print("GEMINI_API_KEY not found or not set. Using placeholder image.")
        return _copy_placeholder_image()

    try:
        genai.configure(api_key=gemini_api_key)

        # --- Placeholder for actual Gemini Text-to-Image API call ---
        # The `google-generativeai` library's direct text-to-image capabilities are specific
        # and might require a particular model name not generally available or through a different method.
        # For example, 'gemini-pro-vision' is for understanding images, not generating them from text.
        # Google's Imagen models are typically used for text-to-image.

        print("NOTE: Attempting conceptual Gemini text-to-image generation.")
        print(f"Intended prompt for Gemini: {prompt_text}")

        # model = genai.GenerativeModel('gemini-1.0-pro-vision-latest') # This is an example, likely not for image gen from text
        # model = genai.GenerativeModel('text-to-image-model-name') # Replace with actual if available
        # response = model.generate_content(prompt_text)
        #
        # If the response contained image bytes:
        # image_bytes = response.parts[0].inline_data.data # Highly speculative path to image bytes
        # img = Image.open(io.BytesIO(image_bytes))
        # img.save(image_abs_save_path, 'PNG')
        # return image_save_subpath
        #
        # If the response contained a URL to an image:
        # image_url = response.candidates[0].content.parts[0].file_data.uri # Highly speculative
        # img_response = requests.get(image_url)
        # img_response.raise_for_status()
        # with open(image_abs_save_path, 'wb') as f:
        #     f.write(img_response.content)
        # return image_save_subpath

        # Since direct, simple text-to-image is not a clear feature of `google-generativeai` for general models:
        raise NotImplementedError("Direct Gemini text-to-image generation via `google-generativeai` is not straightforwardly implemented here. Requires specific models (e.g., Imagen via Vertex AI SDK).")

    except NotImplementedError as nie: # Catching our own NotImplementedError
        print(f"INFO: {nie}")
        print("Falling back to placeholder image logic for Gemini image generation.")
        return _copy_placeholder_image()
    except genai.generation_types.BlockedPromptException as bpe:
        print(f"Gemini API Blocked Prompt Error: {bpe}")
        return _copy_placeholder_image() # Fallback on blocked prompt
    except Exception as e: # Catch other potential google.generativeai errors or general errors
        print(f"An unexpected error occurred with Gemini API: {type(e).__name__} - {e}")
        print("Falling back to placeholder image logic.")
        return _copy_placeholder_image()

def save_summary_to_files(summary_text, base_filename_prefix="summary"):
    """Saves the summary text to Markdown and DOCX files."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_prefix = os.path.splitext(base_filename_prefix)[0]
    base_filename = f"{safe_prefix}_{timestamp}"

    # Store both full path (for logging/internal use) and just filename (for downloads)
    saved_file_details = {'md': None, 'docx': None}

    # Save as Markdown
    md_filename_only = f"{base_filename}.md"
    md_filepath_full = os.path.join(DOCS_FOLDER, md_filename_only)
    try:
        with open(md_filepath_full, 'w', encoding='utf-8') as f:
            f.write(summary_text)
        saved_file_details['md'] = {'full_path': md_filepath_full, 'filename': md_filename_only}
        print(f"Summary saved to {md_filepath_full}")
    except Exception as e:
        print(f"Error saving Markdown file: {e}")

    # Save as DOCX
    docx_filename_only = f"{base_filename}.docx"
    docx_filepath_full = os.path.join(DOCS_FOLDER, docx_filename_only)
    try:
        doc = Document()
        # Using the new parser for potentially structured summary
        parsed_summary_for_docx = parse_structured_summary(summary_text)
        if "Full Summary" in parsed_summary_for_docx and len(parsed_summary_for_docx.keys()) == 1:
            # If it's just a single block of text
            doc.add_paragraph(parsed_summary_for_docx["Full Summary"])
        else: # Attempt to add with headings
            for section, content in parsed_summary_for_docx.items():
                if section != "Error": # Skip error messages
                    doc.add_heading(section, level=1)
                    doc.add_paragraph(content)

        if not doc.paragraphs and not doc.sections: # If nothing was added (e.g. empty summary_text)
             doc.add_paragraph(summary_text) # Fallback to raw text

        doc.save(docx_filepath_full)
        saved_file_details['docx'] = {'full_path': docx_filepath_full, 'filename': docx_filename_only}
        print(f"Summary saved to {docx_filepath_full}")
    except Exception as e:
        print(f"Error saving DOCX file: {e}")

    return saved_file_details


def summarize_text_with_ai(text_to_summarize):
    """
    Summarizes the given text using an AI model.
    Currently returns a mock summary.
    Uses OpenAI API key from app.config.
    """
    openai_api_key = app.config.get('OPENAI_API_KEY')
    if not openai_api_key:
        print("OPENAI_API_KEY not found in environment variables or app.config. Returning mock summary.")
        return """Abstract: OpenAI API key not found. This is a mock abstract.
Introduction: The system requires an OpenAI API key for real summarization with GPT.
Results: Mock results are shown.
Discussion: Please set the OPENAI_API_KEY environment variable for actual AI processing.
"""

    try:
        client = openai.OpenAI(api_key=openai_api_key)

        prompt_text = f"""
Please summarize the following research paper text.
Structure the summary into these distinct sections: Abstract, Introduction, Results, and Discussion.
Ensure each section is clearly labeled (e.g., "Abstract: ...", "Introduction: ...").

Research Paper Text:
{text_to_summarize}
"""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo", # Or another suitable model like gpt-4-turbo-preview
            messages=[
                {"role": "system", "content": "You are a helpful assistant skilled in summarizing research papers into structured formats."},
                {"role": "user", "content": prompt_text}
            ],
            temperature=0.5, # Adjust for creativity vs. factuality
            # max_tokens can be set if needed, but we want a full summary
        )

        if response.choices and response.choices[0].message and response.choices[0].message.content:
            structured_summary = response.choices[0].message.content.strip()
            # The summary should ideally already be in the "Section: Content" format.
            # The existing parse_structured_summary function in the /results route will handle it.
            return structured_summary
        else:
            print("OpenAI API returned an empty or invalid response.")
            return "Error: OpenAI API returned an empty response. Please check the API status or try again."

    except openai.APIConnectionError as e:
        print(f"OpenAI API Connection Error: {e}")
        return f"Error: OpenAI API Connection Error - {e}"
    except openai.RateLimitError as e:
        print(f"OpenAI API Rate Limit Exceeded: {e}")
        return f"Error: OpenAI API Rate Limit Exceeded - {e}"
    except openai.APIStatusError as e:
        print(f"OpenAI API Status Error: {e}")
        return f"Error: OpenAI API Status Error - {e}"
    except Exception as e:
        print(f"An unexpected error occurred with OpenAI: {e}")
        return f"Error: An unexpected error occurred with the OpenAI API - {e}"

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return redirect(request.url)
        file = request.files['pdf_file']
        if file.filename == '':
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            session['original_filename'] = filename
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            try:
                doc = fitz.open(filepath)
                extracted_text = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    extracted_text += page.get_text("text")
                doc.close()
                session['extracted_text'] = extracted_text
                return redirect(url_for('process_and_summarize')) # Changed from display_summary to results
            except Exception as e:
                print(f"Error extracting text: {e}")
                return redirect(request.url)
        else:
            return redirect(request.url)
    return render_template('index.html')

@app.route('/process_and_summarize')
def process_and_summarize():
    extracted_text = session.pop('extracted_text', None)
    if not extracted_text:
        return redirect(url_for('upload_file'))

    structured_summary_text = summarize_text_with_ai(extracted_text)
    session['structured_summary_text'] = structured_summary_text # Store raw text for parsing later

    original_pdf_filename = session.get('original_filename', 'uploaded_pdf')
    summary_file_prefix = os.path.splitext(original_pdf_filename)[0]

    # save_summary_to_files now returns a dict with 'full_path' and 'filename' for each type
    saved_file_details = save_summary_to_files(structured_summary_text, base_filename_prefix=summary_file_prefix)
    session['saved_summary_details'] = saved_file_details # e.g. {'md': {'full_path': '...', 'filename': '...'}, ...}

    # Generate and store visualization (placeholder)
    # Pass the structured_summary_text (string from OpenAI) to the Anthropic prompter
    vis_prompt = generate_visualization_prompt_with_anthropic(structured_summary_text)
    session['visualization_prompt'] = vis_prompt

    image_filename_prefix = os.path.splitext(original_pdf_filename)[0]
    generated_image_path = generate_image_with_ai(vis_prompt, base_filename_prefix=image_filename_prefix)
    session['visualization_image_path'] = generated_image_path

    return redirect(url_for('results')) # Changed from display_summary to results

@app.route('/results') # Renamed from display_summary
def results():
    summary_text = session.get('structured_summary_text', 'No summary generated.')
    saved_details = session.get('saved_summary_details', {}) # e.g. {'md': {'filename': '...'}, ...}
    vis_prompt = session.get('visualization_prompt', 'No visualization prompt generated.')
    image_path = session.get('visualization_image_path', None)

    parsed_summary = parse_structured_summary(summary_text)

    md_fn = saved_details.get('md', {}).get('filename') if saved_details.get('md') else None
    docx_fn = saved_details.get('docx', {}).get('filename') if saved_details.get('docx') else None

    # For display in template, could also pass full paths if needed for messages
    # md_full_path = saved_details.get('md', {}).get('full_path')
    # docx_full_path = saved_details.get('docx', {}).get('full_path')


    return render_template('result.html',
                           summary_data=parsed_summary,
                           visualization_prompt=vis_prompt,
                           visualization_image_path=image_path,
                           md_filename=md_fn,
                           docx_filename=docx_fn)

@app.route('/download_doc/<path:filename>')
def download_doc(filename):
    return send_from_directory(DOCS_DIR, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
